import tkinter as tk
from tkinter import ttk, filedialog
import requests
import socket
import time
import threading
from queue import Queue
from docx import Document

def is_valid_ip(ip):
    try:
        socket.inet_aton(ip)
        return True
    except socket.error:
        return False

def scan_ip():
    target = ip_entry.get()
    if is_valid_ip(target):
        t_IP = target
    else:
        t_IP = socket.gethostbyname(target)
        
    try:
        payload = {'key': 'EDCB16195A4ADD175246575E74F0FC77', 'ip': t_IP, 'format': 'json'}
        api_result = requests.get('https://api.ip2location.io/', params=payload)
        if api_result.status_code == 200:
            data = api_result.json()
            ip = data.get("ip")
            country = data.get("country_code")
            country_name = data.get("country_name")
            region_name = data.get("region_name")
            city_name = data.get("city_name")
            latitude = data.get("latitude")
            longitude = data.get("longitude")
            time_zone = data.get("time_zone")
            asn = data.get("asn")
            as1 = data.get("as")
            
            result_tree.delete(*result_tree.get_children())
            result_tree.insert("", "end", values=("IP:", ip))
            result_tree.insert("", "end", values=("Country:", country))
            result_tree.insert("", "end", values=("Country Name:", country_name))
            result_tree.insert("", "end", values=("Region Name:", region_name))
            result_tree.insert("", "end", values=("City Name:", city_name))
            result_tree.insert("", "end", values=("Latitude:", latitude))
            result_tree.insert("", "end", values=("Longitude:", longitude))
            result_tree.insert("", "end", values=("Time Zone:", time_zone))
            result_tree.insert("", "end", values=("ASN:", asn))
            result_tree.insert("", "end", values=("AS:", as1))
            result_label.config(text="IP adress ma'lumotlari olindi", fg="green")  

    except Exception as e:
        result_label.config(text="Error: " + str(e), fg="red")

def download_docx():
    target = ip_entry.get()
    if is_valid_ip(target):
        t_IP = target
    else:
        t_IP = socket.gethostbyname(target)
    
    try:
        payload = {'key': 'EDCB16195A4ADD175246575E74F0FC77', 'ip': t_IP, 'format': 'json'}
        api_result = requests.get('https://api.ip2location.io/', params=payload)
        if api_result.status_code == 200:
            data = api_result.json()
            ip = data.get("ip")
            country = data.get("country_code")
            country_name = data.get("country_name")
            region_name = data.get("region_name")
            city_name = data.get("city_name")
            latitude = data.get("latitude")
            longitude = data.get("longitude")
            time_zone = data.get("time_zone")
            asn = data.get("asn")
            as1 = data.get("as")

            doc = Document()
            doc.add_heading('IP Address Information', level=1)
            doc.add_paragraph(f"IP: {ip}")
            doc.add_paragraph(f"Country: {country}")
            doc.add_paragraph(f"Country Name: {country_name}")
            doc.add_paragraph(f"Region Name: {region_name}")
            doc.add_paragraph(f"City Name: {city_name}")
            doc.add_paragraph(f"Latitude: {latitude}")
            doc.add_paragraph(f"Longitude: {longitude}")
            doc.add_paragraph(f"Time Zone: {time_zone}")
            doc.add_paragraph(f"ASN: {asn}")
            doc.add_paragraph(f"AS: {as1}")

            file_path = filedialog.asksaveasfilename(defaultextension=".docx")
            if file_path:
                doc.save(file_path)
                result_label.config(text="Muvofaqiyatli saqlandi", fg="green")
            else:
                result_label.config(text="Saqlanmadi ... ", fg="red")

    except Exception as e:
        result_label.config(text="Error: " + str(e), fg="red")

def scan_ports():
    target = ip_entry.get()
    from_port = int(from_entry.get())
    to_port = int(to_entry.get())
    try:
        ip = socket.gethostbyname(target)
    except socket.gaierror:
        result_label.config(text="Hostname topilmadi", fg="red")
        return

    result_label.config(text="Skaner qilindi", fg="blue")
    for port in range(from_port, to_port + 1):
        t = threading.Thread(target=scan_port, args=(ip, port))
        t.start()

def scan_port(ip, port):
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        s.settimeout(0.5)
        result = s.connect_ex((ip, port))
        s.close()
        if result == 0:
            status = "Open"
        else:
            status = "Closed"
        result_tree.insert("", "end", values=(port, status))
    except:
        pass

root = tk.Tk()
root.title("IP, Domain, and Port Scanner")
root.config(bg="light green")

label_font = ('Arial', 12)

ip_frame = tk.Frame(root)
ip_frame.pack(side=tk.TOP, padx=10, pady=10)

ip_label = tk.Label(ip_frame, text="IP address yoki Domain :", font=label_font)
ip_label.pack(side=tk.LEFT, padx=(0, 10))

ip_entry = tk.Entry(ip_frame, font=label_font)
ip_entry.pack(side=tk.LEFT)

port_frame = tk.Frame(root)
port_frame.pack(side=tk.TOP, padx=10, pady=10)

from_label = tk.Label(port_frame, text="Boshlangich Port:", font=label_font)
from_label.pack(side=tk.LEFT)

from_entry = tk.Entry(port_frame, font=label_font)
from_entry.pack(side=tk.LEFT)

to_label = tk.Label(port_frame, text="Oxirgi Port:", font=label_font)
to_label.pack(side=tk.LEFT)

to_entry = tk.Entry(port_frame, font=label_font)
to_entry.pack(side=tk.LEFT)

result_tree = ttk.Treeview(root, columns=('Property', 'Value'), show='headings')
result_tree.heading('Property', text='Property')
result_tree.heading('Value', text='Value')
result_tree.pack()

result_label = tk.Label(root, text='', font=label_font)
result_label.pack(pady=10)

scan_ip_button = tk.Button(root, text="Scan IP", command=scan_ip, font=label_font)
scan_ip_button.pack()

scan_ports_button = tk.Button(root, text="Scan Ports", command=scan_ports, font=label_font)
scan_ports_button.pack()

download_button = tk.Button(root, text="Wordni yuklash", command=download_docx, font=label_font)
download_button.pack()

root.mainloop()
